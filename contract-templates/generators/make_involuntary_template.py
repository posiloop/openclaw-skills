#!/usr/bin/env python3
"""Editable A4 single-page template for 離職協議書 (益循生活) — 以戴鈺玲版條款為內容、黎陸梅版為版型。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Narrow margins to fit on one A4
for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '標楷體')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.25

def set_font(p, size=10.5, bold=False):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rf = rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rPr.append(rf)
        rf.set(qn('w:eastAsia'), '標楷體')
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
tr = title.add_run('離職協議書')
tr.font.size = Pt(20)
tr.font.bold = True
rPr = tr._r.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
rf.set(qn('w:eastAsia'), '標楷體')
rPr.append(rf)

# Parties
p = doc.add_paragraph()
p.add_run('                    益循生活股份有限公司          （以下簡稱甲方）')
set_font(p)
p = doc.add_paragraph()
p.add_run('立協議書人           ＿＿＿＿＿＿＿＿＿         （以下簡稱乙方）')
set_font(p)

# Preamble
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('乙方自民國（下同）＿＿年＿＿月＿＿日起於甲方任職，經甲方於＿＿年＿＿月＿＿日依勞動基準法相關規定通知資遣乙方，雙方之全職僱傭合約於＿＿年＿＿月＿＿日（以下稱「離職日」）終止，雙方權益關係約定如下：')
set_font(p)

# Clause 1 + 補充說明 (blanks)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.8)
p.add_run('一、甲方應於確認乙方履行第四條之義務後，給付共計新台幣（下同）＿＿＿＿＿＿＿元整作為乙方之離職金（包括但不限於任職期間之補償、賠償、勞務報酬、獎金、福利、加班費、特別休假未休工資、預告期間工資、資遣費及其他在職期間和本協議約定所生之一切款項）。')
set_font(p)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.2)
p.add_run('補充說明：')
set_font(p, bold=True)

supp = [
    '1. 上述離職金包含：＿月份薪資新臺幣＿＿＿＿＿元、＿月份薪資新臺幣＿＿＿＿＿元、資遣費新臺幣＿＿＿＿＿元，合計新臺幣＿＿＿＿＿＿元。',
    '2. 甲方應於乙方簽署離職協議書當日起，於下個月 10 號前以匯款方式給付予乙方。',
    '3. 甲方依前項約定完成給付後，雙方確認就本件勞雇關係之工資、資遣費及相關補償請求權均已結清，日後互不請求。',
]
for s in supp:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.7)
    p.paragraph_format.space_after = Pt(1)
    p.add_run(s)
    set_font(p)

# Clauses 2–7
rest = [
    ('二、', '乙方保證於離職日後不得再就任職期間之勞動關係，對甲方進行任何法律上之請求或主張，包括但不限於向勞工或其他行政主管機關之檢舉或申訴、提起民事訴訟或刑事告訴、告發。'),
    ('三、', '乙方於甲方任職時所知悉所有資訊（包括但不限於客戶資訊、營運細節及報價）及本協議書全部內容，均應視為機密事項，乙方不得將前述資訊散佈或揭露予任何第三人（包括但不限於甲方員工）知悉，雙方均不得有傳述、散佈對他方、他方人員及他方關係企業不利之言論或行為。'),
    ('四、', '乙方應於離職日前將甲方指定之工作內容完整交接予甲方指定之第三人。'),
    ('五、', '倘一方知悉他方有任何違反本協議之情事，非違約方得請求因此所生之全部損害賠償（包括但不限於商譽損害、罰鍰、律師費、訴訟費用）。'),
    ('六、', '因本協議書引起之爭議，雙方同意以臺灣臺北地方法院作為第一審管轄法院。'),
    ('七、', '本協議書正本一式兩份，由雙方各執乙份為憑。'),
]
for num, text in rest:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(num + text)
    set_font(p)

# Signatures
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.add_run('立協議書人：'); set_font(p, bold=True)

p = doc.add_paragraph(); p.add_run('甲  方：益循生活股份有限公司'); set_font(p)
p = doc.add_paragraph(); p.add_run('代表人：吳哲安'); set_font(p)
p = doc.add_paragraph(); p.add_run('聯絡地址：臺北市文山區羅斯福路 6 段 146 號 15 樓之 7'); set_font(p)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3)
p.add_run('乙  方：＿＿＿＿＿＿＿＿＿'); set_font(p)
p = doc.add_paragraph(); p.add_run('身份證字號：＿＿＿＿＿＿＿＿＿＿'); set_font(p)
p = doc.add_paragraph(); p.add_run('戶籍地址：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿'); set_font(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.add_run('中   華   民   國          年          月          日')
set_font(p)

out = '/Users/chiu/work/yixun-sheng-huo/documents/離職協議書_範本_益循生活.docx'
doc.save(out)
print(f'Saved: {out}')
