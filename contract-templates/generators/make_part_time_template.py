#!/usr/bin/env python3
"""兼職僱傭合約範本 — 益循生活股份有限公司。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '標楷體')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.2

def set_font(p, size=10.5, bold=False):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rf = rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rPr.append(rf)
        rf.set(qn('w:eastAsia'), '標楷體')
        rf.set(qn('w:ascii'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), 'Times New Roman')

def add_para(text, size=10.5, bold=False, indent_cm=0, align=None, space_after=None):
    p = doc.add_paragraph()
    if indent_cm: p.paragraph_format.left_indent = Cm(indent_cm)
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text)
    set_font(p, size=size, bold=bold)
    return p

def add_clause_head(text):
    return add_para(text, size=11.5, bold=True, space_after=2)

# Title
add_para('兼職僱傭合約', size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

# Parties
add_para('甲方（受僱人）：＿＿＿＿＿＿＿＿＿')
add_para('乙方（僱傭人）：益循生活股份有限公司', space_after=4)

add_para('茲為乙方僱用甲方為乙方兼職勞工，訂立本合約如下，以茲共同遵守：', space_after=4)

# 第一條
add_clause_head('第一條')
add_para('甲方到職日：中華民國＿＿年＿＿月＿＿日到職。', indent_cm=0.5)

# 第二條
add_clause_head('第二條')
add_para('職稱：＿＿＿＿＿＿＿＿＿', indent_cm=0.5)
add_para('工時：工作時間依照公司班表執行。', indent_cm=0.5)
add_para('工資：時薪資總額（含稅）新台幣＿＿＿＿＿元整。', indent_cm=0.5)
add_para('工資發給日：按月發給，發給日為每月 10 號，遇假日順延。', indent_cm=0.5)
add_para('工作地點：＿＿＿＿＿＿＿＿＿。乙方因應公司或店面營運或人力調度情形，得調派或指派甲方至其他工作地點給付勞務。', indent_cm=0.5)
add_para('甲方經任用後，若其能力適性與業績不能符合乙方組織功能運作需求，或發生重大事故，或其考核違未達標準，或經營環境變遷，乙方得解除甲方之職務任命，降級或異動。', indent_cm=0.5, space_after=3)

# 第三條 保密義務
add_clause_head('第三條　保密義務')

add_para('1. 機密資訊之定義', bold=True, indent_cm=0.5)
add_para('甲方應於在職期間所取得或知悉之「機密資訊」應採取必要之保密措施，且不得未經乙方事前書面同意而為自己或第三人使用「機密資訊」。', indent_cm=0.8)
add_para('本條之「機密資訊」包含但不限於供應商、甲方之薪資、甲方為履行勞務所知悉乙方、乙方客戶及其關係企業或與任何第三方有關之未公開資訊、乙方營業計劃、業務資料、客戶資料、財務資訊、商業條件、專利技術、工商秘密、營業秘密等其他資訊。', indent_cm=0.8, space_after=3)

add_para('2. 機密資訊之使用與揭露限制', bold=True, indent_cm=0.5)
add_para('(1) 甲方承諾在僱傭期間及離職後，對於所有機密資訊：', indent_cm=0.8)
for s in [
    '・不得為乙方業務以外之任何目的使用',
    '・不得向任何第三方揭露、洩漏或透露',
    '・不得複製、重製或以任何方式記錄（除非執行職務所必需）',
    '・應採取合理之安全措施防止機密資訊之洩露',
]:
    add_para(s, indent_cm=1.2)
add_para('(2) 機密資訊之例外：以下情形不屬於機密資訊：', indent_cm=0.8)
for s in [
    '・已為公眾所知悉之資訊',
    '・甲方能證明係從非保密管道合法取得之資訊',
    '・甲方於受僱前已合法持有之資訊',
    '・依法律、法規或政府機關要求必須揭露之資訊（但甲方應事先通知乙方）',
]:
    add_para(s, indent_cm=1.2)

add_para('3. 前僱主機密資訊之保護', bold=True, indent_cm=0.5)
add_para('甲方承諾不會不當使用、揭露或引導乙方使用任何前僱主或其他第三方之專有資訊或營業秘密。甲方保證不會將任何前僱主或第三方之未公開文件、專有資訊或營業秘密帶入乙方工作場所或技術系統。', indent_cm=0.8, space_after=3)

add_para('4. 第三方機密資訊', bold=True, indent_cm=0.5)
add_para('甲方理解乙方已經或將會從與乙方有業務往來之第三方（如客戶、供應商、授權方、被授權方、合作夥伴等，以下稱「關聯第三方」）收到其機密或專有資訊（以下稱「關聯第三方機密資訊」），乙方對該等資訊負有保密義務且僅能用於特定限定目的。', indent_cm=0.8)
add_para('甲方同意在僱傭期間及離職後：', indent_cm=0.8)
for s in [
    '・對所有關聯第三方機密資訊負有最嚴格之保密義務',
    '・不得使用或向任何個人、公司或第三方揭露關聯第三方機密資訊',
    '・僅能在執行乙方工作且符合乙方與該關聯第三方協議之範圍內使用該資訊',
    '・遵守乙方不時制定之有關關聯第三方及關聯第三方機密資訊之政策及指引',
]:
    add_para(s, indent_cm=1.2)

add_para('5. 智慧財產權之歸屬', bold=True, indent_cm=0.5)
add_para('甲方在僱傭期間內（包括非上班時間），單獨或與他人共同構思、發現、創作、發明、開發或完成之任何可受著作權保護之資料、記錄、圖樣、設計、標誌、發明、改良、開發、發現、構想及營業秘密，以及任何相關之著作權、專利權、營業秘密、遮罩著作權或其他智慧財產權（統稱「發明創作」），其所有權利、所有權及利益均完全歸屬於乙方所有。', indent_cm=0.8)
add_para('甲方同意：', indent_cm=0.8)
for s in [
    '・立即向乙方完整揭露任何發明創作',
    '・將所有發明創作之權利、所有權及利益讓與並移轉予乙方',
    '・本讓與包括對尚未存在之發明創作之現時讓與',
    '・簽署乙方認為必要之文件以確保、維護及執行乙方對發明創作之權利',
]:
    add_para(s, indent_cm=1.2)

add_para('6. 資訊及財產之返還', bold=True, indent_cm=0.5)
add_para('甲方於離職時或乙方要求時，應立即返還乙方所有財產，包括但不限於：', indent_cm=0.8)
for s in [
    '・所有機密資訊及關聯第三方機密資訊',
    '・所有發明創作之有形體現物',
    '・所有電子儲存資訊及存取密碼',
    '・公司信用卡、記錄、資料、筆記、報告、檔案、提案、清單、通訊記錄、規格、圖樣、藍圖、草圖、材料、照片、圖表及任何其他文件及財產',
]:
    add_para(s, indent_cm=1.2)
add_para('甲方不得保留、複製或向他人交付任何上述資訊或財產。', indent_cm=0.8, space_after=3)

add_para('7. 違反保密義務之後果', bold=True, indent_cm=0.5)
add_para('甲方理解並同意，違反本條保密義務將對乙方造成無法彌補之損害。除法律規定之其他救濟外，乙方有權：', indent_cm=0.8)
for s in [
    '・立即終止僱傭關係',
    '・請求法院核發禁制令',
    '・請求損害賠償（包括但不限於實際損失、合理律師費用）',
    '・請求返還因違反保密義務所獲得之一切利益',
]:
    add_para(s, indent_cm=1.2)
add_para('本條規定之保密義務於本契約終止或屆滿後仍繼續有效。', indent_cm=0.8, space_after=3)

# 第四條
add_clause_head('第四條　終止條款')
add_para('非有下列情形之一者，乙方不得預告甲方終止本合約：①歇業或轉讓時。②虧損或業務緊縮時。③不可抗力暫停工作在一個月以上時。④業務性質變更，有減少勞工必要，又無適當工作可供安置時。⑤乙方對於所擔任之工作確不能勝任時。', indent_cm=0.5)
add_para('乙方依前項規定終止本合約者，其預告期間應依下列規定辦理：①在乙方繼續工作三個月以上一年未滿者，於十日前預告之。②在乙方繼續工作一年以上未滿三年者，於廿日前預告之。③在乙方繼續工作三年以上者，於三十日前預告之。', indent_cm=0.5)
add_para('甲方有下列情形之一者，乙方得不經預告終止契約：①於訂立本合約時為虛偽意思表示，使乙方誤信而有受損害之虞者。②對於乙方、乙方家屬、乙方代理人或其他共同工作之勞工及客戶，實施暴行或有重大侮辱之行為者。③受有期徒刑以上刑之宣告確定，而未諭知緩刑或未准易科罰金者。④違反本合約或工作規則，情節重大者。⑤故意損耗機器、工具、原料、產品，或其他乙方所有物品，或故意洩漏乙方技術上、營業上之秘密，致乙方受有損害者。⑥無正當理由繼續曠職三日，或一個月內曠職達六日者。⑦甲方無故洩漏乙方或乙方客戶之工商秘密等應保密事項者。', indent_cm=0.5)
add_para('其餘未盡事宜，悉依乙方工作規則、《勞動基準法》及其他相關法令之規定辦理。', indent_cm=0.5)
add_para('甲方應於本合約終止日前製作交接清單且備齊相關資訊檔案，完成交接予乙方指定之第三人。', indent_cm=0.5, space_after=3)

# 第五條
add_clause_head('第五條　管轄法院')
add_para('本協議應以中華民國法律為準據法及解釋之依據。雙方就本協議之爭議或歧見，應儘先誠意協商解決，如無法以和平方式於合理期間解決，雙方協議以臺灣臺北地方法院為第一審管轄法院。', indent_cm=0.5, space_after=3)

# 第六條
add_clause_head('第六條　其他約定')
add_para('本協議未盡事宜，經甲、乙雙方同意，得以書面另訂補充協議。', indent_cm=0.5)
add_para('本協議一式兩份，雙方各執一份為憑，並自甲、乙雙方簽署後生效。', indent_cm=0.5, space_after=6)

# 簽署頁
add_para('（簽署頁）', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

add_para('受僱人（甲方）：＿＿＿＿＿＿＿＿＿＿＿＿')
add_para('姓　　名：＿＿＿＿＿＿＿＿＿＿＿＿')
add_para('身分證字號：＿＿＿＿＿＿＿＿＿＿＿＿')
add_para('聯絡電話：＿＿＿＿＿＿＿＿＿＿＿＿')
add_para('聯絡地址：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿', space_after=6)

add_para('僱傭人（乙方）')
add_para('名　　稱：益循生活股份有限公司')
add_para('負 責 人：吳哲安')
add_para('統一編號：60640265')
add_para('聯絡地址：臺北市文山區羅斯福路 6 段 146 號 15 樓之 7', space_after=8)

add_para('中華民國　　年　　月　　日', align=WD_ALIGN_PARAGRAPH.CENTER)

out = '/Users/chiu/.claude/skills/contract-templates/templates/yixun-sheng-huo/part-time-contract.docx'
doc.save(out)
print(f'Saved: {out}')
# Also copy to work dir for immediate use
import shutil
work_out = '/Users/chiu/work/yixun-sheng-huo/documents/兼職僱傭合約_範本_益循生活.docx'
shutil.copy(out, work_out)
print(f'Also: {work_out}')
